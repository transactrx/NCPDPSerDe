"""Generate a Word (.docx) reference document from the NCPDP JSON Schemas.

Reads pkg/ncpdp/schemas/ (shared $defs + per-transaction request/response
schemas + examples) and produces a Confluence-importable Word document
(upload via a Confluence page's ... menu > Import Word Document).

Each transaction and each shared segment definition is documented with a
field table plus a full JSON layout skeleton showing the exact JSON shape.

Requires: pip install python-docx

Usage (from the repo root):

    python scripts/gen_schema_docx.py [schemas-dir] [output.docx]

Defaults: schemas-dir = pkg/ncpdp/schemas,
          output      = NCPDP-JSON-Schema-Reference.docx
"""

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SCHEMAS_DIR = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("pkg/ncpdp/schemas")
OUT_PATH = Path(sys.argv[2]) if len(sys.argv) > 2 else Path("NCPDP-JSON-Schema-Reference.docx")

HEADER_BG = "D9E2F3"  # light blue header row shading

# Internal-only properties, populated by the NCPDPSerDe package itself and
# never supplied by callers, so they are omitted from the documentation:
#   - RawValue/Size in the header wrapper ({RawValue, Size, Value})
#   - Raw captures (SegmentId.Raw, Claims[].Raw)
#   - SegmentId (segment identifiers are auto-populated during serialization)
INTERNAL_HEADER_FIELDS = {"RawValue", "Size"}
INTERNAL_FIELDS = {"Raw", "SegmentId"}


def visible_properties(schema):
    """Properties of an object schema, minus internal bookkeeping fields
    (Raw captures, and RawValue/Size alongside a Value property)."""
    props = schema.get("properties", {})
    skip = set(INTERNAL_FIELDS)
    if "Value" in props:
        skip |= INTERNAL_HEADER_FIELDS
    return {k: v for k, v in props.items() if k not in skip}

# Preferred display order for transaction schema files
FILE_ORDER = [
    "billing", "reversal", "rebill", "eligibility",
    "information", "informationRebill", "informationReversal",
    "priorAuthorization", "priorAuthorizationReversal",
    "priorAuthorizationInquiry", "priorAuthorizationRequestOnly",
    "predeterminationOfBenefits",
    "serviceBilling", "serviceRebill", "serviceReversal",
    "controlledSubstanceReporting", "controlledSubstanceReportingRebill",
    "controlledSubstanceReportingReversal",
]


def load_json(path):
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def ordered_files(directory):
    files = {p.stem: p for p in directory.glob("*.json")}
    result = [files.pop(stem) for stem in FILE_ORDER if stem in files]
    result.extend(files.values())  # anything not in the preferred list
    return result


def type_display(schema):
    """Human-readable type for a schema node."""
    if "$ref" in schema:
        return schema["$ref"].rsplit("/", 1)[-1]
    t = schema.get("type")
    if isinstance(t, list):
        base = [x for x in t if x != "null"]
        t = base[0] if base else "null"
    if t == "array":
        items = schema.get("items", {})
        return f"array of {type_display(items)}" if items else "array"
    if "const" in schema:
        return f'constant "{schema["const"]}"'
    return t or "object"


def constraints_display(schema):
    parts = []
    if "maxLength" in schema:
        parts.append(f"max {schema['maxLength']} chars")
    if "pattern" in schema:
        parts.append(f"pattern {schema['pattern']}")
    if "minimum" in schema:
        parts.append(f"min {schema['minimum']}")
    if "maximum" in schema:
        parts.append(f"max {schema['maximum']}")
    if "maxItems" in schema:
        parts.append(f"max {schema['maxItems']} items")
    if "enum" in schema:
        parts.append("one of: " + ", ".join(str(v) for v in schema["enum"]))
    if "const" in schema:
        parts.append(f"always \"{schema['const']}\"")
    return "; ".join(parts)


CODE_RE = re.compile(r"\(([A-Z0-9]{2})\)\s*$")
F6_RE = re.compile(r"\s*-\s*F6\s*$")


def parse_desc(description):
    """Split a description into (NCPDP code, cleaned text, is_f6_only).

    Descriptions follow the convention 'Field Name (C2)' with F6-only fields
    suffixed ' - F6' (e.g. 'Patient Middle Name (0C) - F6')."""
    if not description:
        return "", "", False
    is_f6 = bool(F6_RE.search(description))
    desc = F6_RE.sub("", description)
    m = CODE_RE.search(desc)
    if m:
        return m.group(1), desc[: m.start()].strip(), is_f6
    return "", desc.strip(), is_f6


def split_code(description):
    """Extract trailing NCPDP field code like '(C2)' from a description."""
    code, desc, _ = parse_desc(description)
    return code, desc


def walk_fields(properties, prefix, rows):
    """Flatten nested object properties into
    (path, code, version, type, constraints, desc) rows. The version column
    holds 'F6' for F6-only fields and is blank for fields available in both
    D.0 and F6."""
    for name, sub in properties.items():
        path = f"{prefix}{name}"
        if "$ref" in sub:
            ref = sub["$ref"].rsplit("/", 1)[-1]
            rows.append((path, "", "", ref, "", f"See shared definition: {ref}"))
            continue
        t = sub.get("type")
        if isinstance(t, list):
            t = next((x for x in t if x != "null"), t[0])
        if t == "object" and "properties" in sub:
            walk_fields(visible_properties(sub), path + ".", rows)
        elif t == "array":
            items = sub.get("items", {})
            if "$ref" in items:
                code, desc, is_f6 = parse_desc(sub.get("description", ""))
                ref = items["$ref"].rsplit("/", 1)[-1]
                rows.append((path, code, "F6" if is_f6 else "", f"array of {ref}",
                             constraints_display(sub),
                             desc or f"See shared definition: {ref}"))
            elif items.get("type") == "object" and "properties" in items:
                code, desc, is_f6 = parse_desc(sub.get("description", ""))
                rows.append((path, code, "F6" if is_f6 else "",
                             "array (repeating group)", constraints_display(sub), desc))
                walk_fields(items["properties"], path + "[]. ".rstrip() , rows)
            else:
                code, desc, is_f6 = parse_desc(sub.get("description", ""))
                rows.append((path, code, "F6" if is_f6 else "", type_display(sub),
                             constraints_display(sub), desc))
        else:
            code, desc, is_f6 = parse_desc(sub.get("description", ""))
            rows.append((path, code, "F6" if is_f6 else "", type_display(sub),
                         constraints_display(sub), desc))


def leaf_placeholder(schema):
    """Placeholder value for a leaf field in a JSON layout skeleton,
    e.g. 'string (C2, max 20)' or '\"B1\" (constant)'."""
    if "const" in schema:
        return f'"{schema["const"]}" (constant)'
    t = schema.get("type")
    if isinstance(t, list):
        t = next((x for x in t if x != "null"), t[0])
    parts = []
    code, _, is_f6 = parse_desc(schema.get("description", ""))
    if code:
        parts.append(code)
    if is_f6:
        parts.append("F6 only")
    if "maxLength" in schema:
        parts.append(f"max {schema['maxLength']}")
    if "pattern" in schema:
        parts.append(f"pattern {schema['pattern']}")
    if "minimum" in schema or "maximum" in schema:
        lo, hi = schema.get("minimum", ""), schema.get("maximum", "")
        parts.append(f"{lo}-{hi}")
    if "enum" in schema:
        parts.append("|".join(str(v) for v in schema["enum"]))
    return f"{t or 'object'} ({', '.join(parts)})" if parts else (t or "object")


def json_skeleton(schema, defs, expand_refs=False, _depth=0):
    """Build a JSON-shaped skeleton (dict/list/placeholder strings) for a
    schema node. $refs are expanded inline when expand_refs is True, otherwise
    rendered as a pointer to the shared definition."""
    if "$ref" in schema:
        ref = schema["$ref"].rsplit("/", 1)[-1]
        if expand_refs and ref in defs and _depth < 20:
            return json_skeleton(defs[ref], defs, expand_refs, _depth + 1)
        return f"{{...}} see: {ref}"
    t = schema.get("type")
    if isinstance(t, list):
        t = next((x for x in t if x != "null"), t[0])
    if t == "object" and "properties" in schema:
        return {
            name: json_skeleton(sub, defs, expand_refs, _depth + 1)
            for name, sub in visible_properties(schema).items()
        }
    if t == "array":
        items = schema.get("items", {})
        return [json_skeleton(items, defs, expand_refs, _depth + 1)] if items else []
    return leaf_placeholder(schema)


def add_json_layout(doc, schema, defs, expand_refs=False, heading=None, level=3):
    if heading:
        doc.add_heading(heading, level=level)
    skeleton = json_skeleton(schema, defs, expand_refs)
    add_code_block(doc, json.dumps(skeleton, indent=2))


# ---------------------------------------------------------------- docx helpers

def shade_cell(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table(doc, headers, rows, widths=None, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        shade_cell(hdr[i], HEADER_BG)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(font_size)
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def add_code_block(doc, text, font_size=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(font_size)
    # shade the paragraph background light gray
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def field_table(doc, def_schema):
    rows = []
    walk_fields(visible_properties(def_schema), "", rows)
    if rows:
        add_table(
            doc,
            ["Field", "NCPDP Code", "Ver.", "Type", "Constraints", "Description"],
            rows,
            widths=[1.8, 0.7, 0.5, 1.1, 1.2, 2.1],
        )


def transaction_section(doc, schema, defs):
    doc.add_heading(schema.get("title", "Transaction"), level=2)
    if schema.get("description"):
        doc.add_paragraph(schema["description"])
    required = set(schema.get("required", []))

    # Top-level structure table
    rows = []
    claims_schema = None
    claims_name = None
    for name, sub in schema.get("properties", {}).items():
        req = "Yes" if name in required else "No"
        if "$ref" in sub:
            ref = sub["$ref"].rsplit("/", 1)[-1]
            desc = defs.get(ref, {}).get("description", "")
            rows.append((name, ref, req, desc))
        elif sub.get("type") == "array" and "properties" in sub.get("items", {}):
            claims_schema, claims_name = sub, name
            rows.append((name, "array of claim records", req, sub.get("description", "")))
        elif sub.get("type") == "array" and "$ref" in sub.get("items", {}):
            ref = sub["items"]["$ref"].rsplit("/", 1)[-1]
            rows.append((name, f"array of {ref}", req, sub.get("description", "")))
        else:
            rows.append((name, type_display(sub), req, sub.get("description", "")))
    doc.add_heading("Structure", level=3)
    add_table(doc, ["Property", "Type", "Required", "Description"], rows,
              widths=[1.4, 1.9, 0.7, 3.4])

    # Header field detail
    header = schema.get("properties", {}).get("Header", {})
    value = header.get("properties", {}).get("Value", {})
    if value.get("properties"):
        doc.add_heading("Header Fields (Header.Value)", level=3)
        hreq = set(value.get("required", []))
        rows = []
        for name, sub in value["properties"].items():
            code, desc = split_code(sub.get("description", ""))
            rows.append((name, type_display(sub), "Yes" if name in hreq else "No",
                         constraints_display(sub), desc))
        add_table(doc, ["Field", "Type", "Required", "Constraints", "Description"],
                  rows, widths=[1.9, 0.8, 0.7, 1.6, 2.4])

    # Per-claim segment list
    if claims_schema is not None:
        doc.add_heading(f"Per-Claim Segments ({claims_name}[])", level=3)
        rows = []
        for name, sub in visible_properties(claims_schema["items"]).items():
            if "$ref" in sub:
                ref = sub["$ref"].rsplit("/", 1)[-1]
                rows.append((name, ref, defs.get(ref, {}).get("description", "")))
            elif sub.get("type") == "array" and "$ref" in sub.get("items", {}):
                ref = sub["items"]["$ref"].rsplit("/", 1)[-1]
                rows.append((name, f"array of {ref}", sub.get("description", "")))
            else:
                rows.append((name, type_display(sub), sub.get("description", "")))
        add_table(doc, ["Property", "Type", "Description"], rows,
                  widths=[1.9, 2.2, 3.3])

    # Full JSON layout skeleton (segment refs point to the shared definitions,
    # whose own full layouts appear in the Segment and Type Definitions section)
    add_json_layout(doc, schema, defs, expand_refs=False,
                    heading="JSON Layout", level=3)


# ------------------------------------------------------------------- document

def main():
    shared = load_json(SCHEMAS_DIR / "ncpdp-schemas.json")
    defs = shared["$defs"]

    doc = Document()
    doc.core_properties.title = "NCPDP D.0/F6 JSON Schema Reference"

    # Slightly wider usable page
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    doc.add_heading("NCPDP D.0/F6 JSON Schema Reference", level=0)
    p = doc.add_paragraph()
    p.add_run(shared.get("description", "")).italic = True
    doc.add_paragraph(
        "This document is generated from the JSON Schema definitions in "
        "pkg/ncpdp/schemas of the NCPDPSerDe repository. The schemas describe the "
        "JSON representation of NCPDP Telecommunication Standard D.0 (and F6) "
        "request and response transactions. They can be used to understand the data "
        "structure of NCPDP transactions, validate JSON payloads before "
        "serialization, generate client code, and document API contracts."
    )

    # ---- Overview
    doc.add_heading("Overview", level=1)

    doc.add_heading("Transaction Codes", level=2)
    add_table(doc, ["Code", "Transaction Type"], [
        ("B1", "Billing (New Rx)"), ("B2", "Reversal"), ("B3", "Rebill"),
        ("E1", "Eligibility Verification"),
        ("N1", "Information Reporting"), ("N2", "Information Reporting Rebill"),
        ("N3", "Information Reporting Reversal"),
        ("P1", "Prior Authorization Request & Billing"),
        ("P2", "Prior Authorization Reversal"), ("P3", "Prior Authorization Inquiry"),
        ("P4", "Prior Authorization Request Only"),
        ("D1", "Predetermination of Benefits"),
        ("S1", "Service Billing"), ("S2", "Service Rebill"), ("S3", "Service Reversal"),
        ("C1", "Controlled Substance Reporting"),
        ("C2", "Controlled Substance Reporting Rebill"),
        ("C3", "Controlled Substance Reporting Reversal"),
    ], widths=[0.8, 4.5])

    doc.add_heading("Key Concepts", level=2)
    concepts = [
        ("Versions (D.0 / F6)", "The schemas cover both NCPDP Telecommunication "
                   "Standard D.0 and F6; the version is auto-detected from the "
                   "header. Every D.0 field is unchanged in F6 (full backward "
                   "compatibility). Fields and segments that only exist in F6 are "
                   "marked 'F6' in the Ver. column of the field tables and in the "
                   "JSON layouts — they are optional and simply omitted for D.0. "
                   "Constraints show the largest value across versions (e.g. the "
                   "BIN/IIN field is 6 characters in D.0 and 8 in F6, so max 8)."),
        ("Header", "Every transaction has a Header containing the BIN (D.0, 6 "
                   "characters) or IIN (F6, 8 characters) in the Bin property, "
                   "version (D0 or F6), transaction code, record count, service "
                   "provider ID, and date of service (CCYYMMDD). Header fields are "
                   "supplied under Header.Value; the RawValue and Size header "
                   "properties and the Raw capture fields (on segments and claim "
                   "records) that appear in the schema files are internal to the "
                   "NCPDPSerDe package and must not be supplied when submitting "
                   "JSON (they are omitted from this document)."),
        ("Claims per transaction", "D.0 allows 1–4 claim records per transmission "
                   "(the Claims array and header RecordCount). F6 allows exactly "
                   "one claim per transmission — the serializer rejects F6 "
                   "transactions with more than one claim group."),
        ("Segments", "Transactions consist of segments identified by AM codes "
                     "(AM01 Patient, AM04 Insurance, AM07 Claim, AM11 Pricing, etc.). "
                     "Each segment is documented in the Segment Definitions section. "
                     "The SegmentId property in the schema files does not need to be "
                     "supplied — NCPDPSerDe fills in each segment's identifier "
                     "automatically during serialization."),
        ("Field codes", "Each field within a segment has a 2-character NCPDP field "
                        "code (e.g. C2 = Cardholder ID, D2 = Rx Number), shown in the "
                        "NCPDP Code column of the field tables."),
        ("Nullable fields", "Most fields are nullable and may be omitted or set to "
                            "null. Only fields marked Required must be provided."),
        ("Counter fields", "Repeating-group count fields (e.g. Coordination of "
                           "Benefits Count 4C, Reject Code Count FA) and per-item "
                           "counters (e.g. DUR/PPS Code Counter 7E) are derived "
                           "automatically during serialization: counts are set to the "
                           "array length (correcting any supplied value) and item "
                           "counters are numbered 1..N. They only need to be supplied "
                           "when using D0 scalar fields instead of the repeating "
                           "array (e.g. Patient ID Count with the single "
                           "IdQualifier/Id fields)."),
        ("DynamicFields / DynamicSegments", "Each segment includes a DynamicFields "
                                            "array (and each transaction a DynamicSegments array) for "
                                            "non-standard or vendor-specific data."),
        ("Dates and times", "Dates use CCYYMMDD (e.g. 20240115); times use HHmm or "
                            "HHmmss."),
        ("Monetary amounts", "Monetary amounts are numbers with 2 decimal places and "
                             "may be positive or negative (overpunch representation in raw "
                             "NCPDP format is handled internally)."),
    ]
    for term, text in concepts:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(term + ": ").bold = True
        p.add_run(text)

    # ---- Request transactions
    doc.add_heading("Request Transactions", level=1)
    for path in ordered_files(SCHEMAS_DIR / "request"):
        transaction_section(doc, load_json(path), defs)

    # ---- Response transactions
    doc.add_heading("Response Transactions", level=1)
    for path in ordered_files(SCHEMAS_DIR / "response"):
        transaction_section(doc, load_json(path), defs)

    # ---- Shared definitions
    doc.add_heading("Segment and Type Definitions", level=1)
    doc.add_paragraph(
        "Shared definitions referenced by the transaction schemas above "
        "(from ncpdp-schemas.json). Field paths use dot notation for nested "
        "objects and [] for repeating groups."
    )

    groups = [
        ("Common Types", [k for k in defs if "." not in k]),
        ("Request Segments", [k for k in defs if k.startswith("request.")]),
        ("Response Segments", [k for k in defs if k.startswith("response.")]),
    ]
    for group_title, keys in groups:
        doc.add_heading(group_title, level=2)
        for key in keys:
            d = defs[key]
            title = d.get("description") or key
            doc.add_heading(f"{key} — {title}" if title != key else key, level=3)
            field_table(doc, d)
            p = doc.add_paragraph()
            p.add_run("JSON layout:").bold = True
            add_json_layout(doc, d, defs, expand_refs=True)

    # ---- Examples
    doc.add_heading("Examples", level=1)
    for path in sorted((SCHEMAS_DIR / "examples").glob("*.json")):
        doc.add_heading(path.stem, level=2)
        add_code_block(doc, json.dumps(load_json(path), indent=2))

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
